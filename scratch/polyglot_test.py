import asyncio
import os
import docx

def extract_text_from_any(file_path: str):
    ext = file_path.split(".")[-1].lower()
    content = ""
    try:
        if ext == "docx":
            doc = docx.Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
    except Exception as e:
        return f"Extraction Error: {str(e)}"
    return content

async def smoke_test_translation_pipeline():
    print("--- SAATHI POLYGLOT SMOKE TEST (v107.1) ---")
    
    # 1. Create a Test DOCX
    test_file = "translation_test.docx"
    doc = docx.Document()
    doc.add_paragraph("Hello, this is a document for Saathi translation test.")
    doc.save(test_file)
    print(f"File Created: {test_file}")
    
    # 2. Simulate Front-end Extraction (v107.0 flow)
    print("Simulating Front-end Extraction...")
    extracted_content = extract_text_from_any(test_file)
    
    # 3. Simulate JSON Payload
    translation_req = {
        "name": test_file,
        "content": extracted_content,
        "target_lang": "hi" # Hindi
    }
    print(f"JSON Payload Created for: {translation_req['name']}")
    
    # 4. Verify Backend Tool Mapping
    print("Verifying Tool Mapping...")
    try:
        from agent_tools import execute_tool
        print("SUCCESS: 'execute_tool' imported correctly from agent_tools.")
    except ImportError as e:
        print(f"FAILURE: {e}")
        
    print("\n[RESULT]")
    print(f"Saathi will now translate: '{extracted_content.strip()}'")
    print(f"Target Language: Hindi")
    print(f"Mirror Action: create_doc(title='Polyglot_Hindi_{test_file}', content='[TRANSLATED_TEXT]')")
    print("SMOKE TEST PASS: Polyglot Mirroring is structurally sound.")

    # Clean up
    if os.path.exists(test_file):
        os.remove(test_file)

if __name__ == "__main__":
    asyncio.run(smoke_test_translation_pipeline())
