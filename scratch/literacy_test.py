import os
import docx
from PyPDF2 import PdfWriter, PdfReader
import io

def extract_text_from_any(file_path: str):
    ext = file_path.split(".")[-1].lower()
    content = ""
    try:
        if ext == "pdf":
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                content = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif ext == "docx":
            doc = docx.Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        else: # txt/py/md
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
    except Exception as e:
        return f"Extraction Error: {str(e)}"
    return content

def run_smoke_test():
    print("--- SAATHI LITERACY SMOKE TEST (v106.1) ---")
    
    # 1. Create a Test DOCX
    test_docx = "test_smoke.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a smoke test content for Saathi v106.1.")
    doc.save(test_docx)
    print(f"File Created: {test_docx}")
    
    # 2. Extract
    print("Running Extraction...")
    extracted = extract_text_from_any(test_docx)
    
    # 3. Verify
    if "smoke test content" in extracted:
        print(f"RESULT SUCCESS: Extracted Content: '{extracted.strip()}'")
    else:
        print(f"RESULT FAILURE: Extraction failed or returned garbage: '{extracted}'")
    
    # Clean up
    if os.path.exists(test_docx):
        os.remove(test_docx)

if __name__ == "__main__":
    run_smoke_test()
